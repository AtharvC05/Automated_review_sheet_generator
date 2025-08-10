import mysql.connector
from docx import Document
import os
import subprocess

def connect_db():
    """Connect to the MySQL database."""
    try:
        return mysql.connector.connect(
            host="localhost",
            user="root",
            password="1234",
            database="project_review"
        )
    except mysql.connector.Error as err:
        raise Exception(f"Database connection error: {err}")

def fetch_project_details(group_id):
    """Fetch project and members' details from the database by group_id."""
    conn = connect_db()
    if conn is None:
        raise Exception("Database connection failed")

    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("SELECT * FROM projects WHERE group_id = %s", (group_id,))
        project = cursor.fetchone()

        cursor.execute("SELECT roll_no, student_name, contact_details FROM members WHERE group_id = %s", (group_id,))
        members = cursor.fetchall()

        if project:
            return {
                "group_id": project["group_id"],
                "project_title": project["project_title"] or "",
                "guide_name": project["guide_name"] or "",
                "mentor_name": project["mentor_name"] or "",
                "mentor_email": project["mentor_email"] or "",
                "mentor_mobile": project["mentor_mobile"] or "",
                "members": [
                    {
                        "roll_no": m["roll_no"] or "",
                        "student_name": m["student_name"] or "",
                        "contact_details": m["contact_details"] or ""
                    } for m in members
                ]
            }
        raise Exception("Group ID not found")
    except mysql.connector.Error as err:
        raise Exception(f"Database query error: {err}")
    finally:
        cursor.close()
        conn.close()

def replace_placeholders(doc, placeholders):
    """Replace placeholders in the Word document."""
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

def convert_to_pdf_libreoffice(docx_path, output_dir):
    """Convert DOCX to PDF using LibreOffice in Ubuntu."""
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ], check=True)
    except subprocess.CalledProcessError as e:
        raise Exception(f"LibreOffice PDF conversion failed: {str(e)}")

def generate_review4_pdf(data):
    """Generate Review-IV PDF and return the file path."""
    group_id = data.get('group_id')
    project_data = fetch_project_details(group_id)
    if not project_data:
        raise Exception("Group ID not found or database error")

    template_path = os.path.abspath('Review-IV Sheet.docx')
    if not os.path.exists(template_path):
        raise Exception("Template file not found")

    doc = Document(template_path)

    placeholders = {
        "{{group_id}}": project_data["group_id"],
        "{{date}}": data.get('date', ''),
        "{{project_title}}": project_data["project_title"],
        "{{guide_name}}": project_data["guide_name"],
        "{{mentor_name}}": project_data["mentor_name"],
        "{{mentor_email}}": project_data["mentor_email"],
        "{{mentor_mobile}}": project_data["mentor_mobile"],
        "{{4.1.1id}}": data.get('que_4.1.1', ''),
        "{{4.1.2id}}": data.get('que_4.1.2', ''),
        "{{4.1.3id}}": data.get('que_4.1.3', ''),
        "{{4.1.4id}}": data.get('que_4.1.4', ''),
        "{{4.1.5id}}": data.get('que_4.1.5', ''),
        "{{4.1.6d}}": data.get('que_4.1.6', ''),
        "{{4.1.1}}": data.get('f4.1.1', ''),
        "{{4.2.1}}": data.get('f4.2.1', ''),
        "{{4.3.1}}": data.get('f4.3.1', ''),
        "{{4.4.1}}": data.get('f4.4.1', ''),
        "{{4.1.2}}": data.get('f4.1.2', ''),
        "{{4.2.2}}": data.get('f4.2.2', ''),
        "{{4.3.2}}": data.get('f4.3.2', ''),
        "{{4.4.2}}": data.get('f4.4.2', ''),
        "{{4.1.3}}": data.get('f4.1.3', ''),
        "{{4.2.3}}": data.get('f4.2.3', ''),
        "{{4.3.3}}": data.get('f4.3.3', ''),
        "{{4.4.3}}": data.get('f4.4.3', ''),
        "{{4.1.4}}": data.get('f4.1.4', ''),
        "{{4.2.4}}": data.get('f4.2.4', ''),
        "{{4.3.4}}": data.get('f4.3.4', ''),
        "{{4.4.4}}": data.get('f4.4.4', ''),
        "{{4.1.5}}": data.get('f4.1.5', ''),
        "{{4.2.5}}": data.get('f4.2.5', ''),
        "{{4.3.5}}": data.get('f4.3.5', ''),
        "{{4.4.5}}": data.get('f4.4.5', ''),
        "{{4.1.s1}}": data.get('f4.1.s1', ''),
        "{{4.2.s1}}": data.get('f4.2.s1', ''),
        "{{4.3.s1}}": data.get('f4.3.s1', ''),
        "{{4.4.s1}}": data.get('f4.4.s1', ''),
        "{{4.c}}": data.get('c4', ''),
    }

    for i, member in enumerate(project_data["members"], start=1):
        placeholders[f"{{{{roll_{i}}}}}"] = member["roll_no"]
        placeholders[f"{{{{student_{i}}}}}"] = member["student_name"]
        placeholders[f"{{{{contact_{i}}}}}"] = member["contact_details"]

    for i in range(len(project_data["members"]) + 1, 5):
        placeholders[f"{{{{roll_{i}}}}}"] = ""
        placeholders[f"{{{{student_{i}}}}}"] = ""
        placeholders[f"{{{{contact_{i}}}}}"] = ""

    replace_placeholders(doc, placeholders)

    filled_doc_path = os.path.abspath('Filled_Form_Review_IV.docx')
    doc.save(filled_doc_path)

    # Convert DOCX to PDF using LibreOffice
    output_dir = os.path.dirname(filled_doc_path)
    convert_to_pdf_libreoffice(filled_doc_path, output_dir)

    pdf_output_path = os.path.join(output_dir, 'Filled_Form_Review_IV.pdf')
    if not os.path.exists(pdf_output_path):
        raise Exception("PDF generation failed")

    return pdf_output_path
